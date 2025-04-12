# 워드 작성 및 PDF 변환 관련 패키지 import
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook
from docx2pdf import convert

# 기존 양식 불러오기
doc = Document('교육 수료증.docx')

# 스타일 적용하기 (글꼴, 크기)
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 수강생 명단 엑셀 파일 불러오기
wb = load_workbook("수강생 명단.xlsx")
ws = wb.active

# 엑셀 명단 참고하여 워드 문서 수정 (반복문)
for i in range(ws.max_row): 
    name = ws.cell(row=i+1, column=1).value
    birth = ws.cell(row=i+1, column=2).value

    if i != 0:
        # 성명 수정
        p = doc.paragraphs[3]
        p.style = style
        p.text = '성    명: ' + name


        # 생년월일 수정
        p = doc.paragraphs[4]
        p.style = style
        p.text = '생년월일: ' + str(birth)[:10]

        #저장
        word_filename = f'교육 수료증_{name}.docx'
        doc.save(word_filename)

        # Word 파일을 PDF로 변환
        pdf_filename = f'교육 수료증_{name}.pdf'
        convert(word_filename, pdf_filename)

#관련 모듈, 클래스 불러오기
import os
import smtplib
from email.header import Header
from email.mime.base import MIMEBase # 메일 내용과 첨부파일을 담는 포맷 클래스
from email.mime.text import MIMEText # 메일 내용 작성 관련 클래스
from email.mime.application import MIMEApplication # 파일 첨부 관련 클래스

#각 메일의 SMTP 서버를 dictionary로 정의
smtp_info = {
    'gmail.com': ('smtp.gmail.com', 587),
    'naver.com': ('smtp.naver.com', 587),
    'outlook.com': ('smtp-mail.outlook.com', 587),
}

#메일 보내는 함수 정의 (발신 메일, 수신 메일, 제목, 본문, 첨부파일 경로, 비밀번호)
def send_email(sender_email, receiver_email, subject, message, attachments=(), password='', subtype='plain'):

    # 멀티 파트 포맷을 객체화
    mail_format = MIMEBase('multipart', 'mixed')

    # 입력한 이메일 주소, 제목, 본문 등을 암호화하여 메일 형식으로 입력
    mail_format['From'] = sender_email #발신 메일을 포맷에 입력
    mail_format['To'] = receiver_email #수신 메일을 포맷에 입력
    mail_format['Subject'] = Header(subject.encode('utf-8'), 'utf-8') #utf-8로 인코딩 후, Header 모듈로 메일 제목 입력
    message = MIMEText(message.encode('utf-8'), _subtype=subtype, _charset='utf-8') #메일 본문 인코딩
    mail_format.attach(message) #인코딩한 본문 mail_format에 입력

    #여러개의 파일을 하나씩 첨부
    for file_path in attachments:
        folder, file = os.path.split(file_path)

        with open(file_path, 'rb') as file_obj: # 첨부 파일 열기
            attachment_contents = file_obj.read()

        attachment = MIMEApplication(attachment_contents, _subtype=subtype)
        attachment.add_header('Content-Disposition', 'attachment', filename=(Header(file, 'utf-8').encode()))
        mail_format.attach(attachment)

    #SMTP 서버 로그인 및 작성된 메일 보내기
    username, host = sender_email.rsplit("@",1) #발신인 메일 주소의 @를 기준으로 username과 host로 나눔

    smtp_server, port = smtp_info[host] #step2의 dict를 이용해서 host와 port 정보들을 받아옴

    # SMTP 서버 접속 여부 확인
    if port == 587:
        smtp = smtplib.SMTP(smtp_server, port)
        rcode1, _ = smtp.ehlo()
        rcode2, _ = smtp.starttls()


    else:
        smtp = smtplib.SMTP_SSL(smtp_server, port)
        rcode1, _ = smtp.ehlo()
        rcode2 = 220

    if rcode1 != 250 or rcode2 != 220:
        smtp.quit()
        return '연결에 실패하였습니다.'

    smtp.login(sender_email, password)
    smtp.sendmail(sender_email, receiver_email, mail_format.as_string())
    smtp.quit()

#실제 함수 실행 부분
sender_address = '보내는 메일 주소'
password = '보내는 메일 비밀번호'

for i in range(ws.max_row):
    name = ws.cell(row=i+2, column=1).value
    attachment_file = [f'교육 수료증_{name}.pdf']
    course = ws.cell(row=i+2, column=3).value
    receiver_address = ws.cell(row=i+2, column=5).value

    subject = f'{course} 수료증 전달'

    message = f"""
    안녕하세요, {name}님.

    몰래컴퍼니 교육 담당자입니다.

    {course} 과정의 수강 완료를 축하드립니다.

    수료증을 첨부파일로 보내드립니다.

    감사합니다.
    """

    try:
        send_email(sender_address, receiver_address, subject, message, attachment_file, password=password)
    except:
        pass