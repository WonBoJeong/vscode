from docx import Document

doc = Document()

doc.add_heading('사장님 몰래하는 파이썬 업무자동화', level = 0)

p = doc.add_paragraph('여러분들의 공부를 응원합니다!')
p.add_run(' 이번시간엔 기존문서를 인덱싱하는 법을 공부해봅시다.').bold = True

doc.add_paragraph('문장 추가 1')
doc.add_paragraph('문장 추가 2')
doc.add_paragraph('문장 추가 3')
doc.add_paragraph('문장 추가 4')

records = (
    (1, '하나', 'one'),
    (2, '둘', 'two'),
    (3, '셋', 'three')
)

table = doc.add_table(rows=1, cols=3)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid'로 설정
table.style = doc.styles['Table Grid']

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No.'
hdr_cells[1].text = '한국어'
hdr_cells[2].text = '영어'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

doc.save('인덱싱 예제 문서.docx')
