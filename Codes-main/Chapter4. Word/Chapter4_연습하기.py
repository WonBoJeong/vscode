from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook

doc = Document('교육 수료증.docx')

#스타일 적용하기 (일반 텍스트, 글꼴은 ‘맑은고딕체’)
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

wb = load_workbook("수강생 명단.xlsx")
ws = wb.active

for i in range(ws.max_row):
    name = ws.cell(row=i+1, column=1).value
    birth = ws.cell(row=i+1, column=2).value

    if i != 0:
        # 성명 수정
        p = doc.paragraphs[3]
        p.text = '성    명: ' + name

        # 생년월일 수정
        p = doc.paragraphs[4]
        p.text = '생년월일: ' + str(birth)[:10]

        #저장
        doc.save('교육 수료증_' + name + '.docx')